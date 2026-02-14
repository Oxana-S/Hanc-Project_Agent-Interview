"""
Document Parser - unified parser for PDF, DOCX, MD, XLSX files.

v1.0: Initial implementation
"""

import re
from pathlib import Path
from typing import List, Optional

import structlog

from .models import DocumentChunk, ParsedDocument

logger = structlog.get_logger("documents")


class DocumentParser:
    """
    Унифицированный парсер документов.

    Поддерживает форматы: PDF, DOCX, MD, XLSX
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".xlsx", ".xls", ".txt"}

    def __init__(self):
        """Инициализация парсера."""
        self._check_dependencies()

    def _check_dependencies(self):
        """Проверить наличие необходимых библиотек."""
        self._has_pymupdf = False
        self._has_docx = False
        self._has_openpyxl = False

        try:
            import fitz  # PyMuPDF

            self._has_pymupdf = True
        except ImportError:
            logger.warning("PyMuPDF not installed, PDF parsing disabled")

        try:
            import docx

            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not installed, DOCX parsing disabled")

        try:
            import openpyxl

            self._has_openpyxl = True
        except ImportError:
            logger.warning("openpyxl not installed, XLSX parsing disabled")

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

    def parse(self, file_path: Path) -> Optional[ParsedDocument]:
        """
        Распарсить документ.

        Args:
            file_path: Путь к файлу

        Returns:
            ParsedDocument или None если формат не поддерживается
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error("File not found", path=str(file_path))
            return None

        # R3-5: Prevent OOM from oversized uploads
        try:
            file_size = file_path.stat().st_size
            if file_size > self.MAX_FILE_SIZE:
                logger.error(
                    "File too large",
                    path=str(file_path),
                    size_mb=round(file_size / (1024 * 1024), 1),
                    max_mb=self.MAX_FILE_SIZE // (1024 * 1024),
                )
                return None
        except OSError as e:
            logger.error("Cannot stat file", path=str(file_path), error=str(e))
            return None

        ext = file_path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.warning("Unsupported file format", extension=ext)
            return None

        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext == ".docx":
                return self._parse_docx(file_path)
            elif ext == ".md" or ext == ".txt":
                return self._parse_markdown(file_path)
            elif ext in (".xlsx", ".xls"):
                return self._parse_xlsx(file_path)
        except Exception as e:
            logger.error("Failed to parse document", path=str(file_path), error=str(e))
            return None

        return None

    def _parse_pdf(self, file_path: Path) -> Optional[ParsedDocument]:
        """Парсинг PDF файла."""
        if not self._has_pymupdf:
            logger.error("PyMuPDF not available")
            return None

        import fitz

        chunks = []
        metadata = {}

        with fitz.open(file_path) as doc:
            metadata["pages"] = len(doc)
            metadata["title"] = doc.metadata.get("title", "")
            metadata["author"] = doc.metadata.get("author", "")

            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                if text.strip():
                    chunks.append(
                        DocumentChunk(
                            content=text.strip(),
                            doc_type="pdf",
                            page=page_num,
                        )
                    )

        logger.info("PDF parsed", path=str(file_path), pages=len(chunks))

        return ParsedDocument(
            filename=file_path.name,
            doc_type="pdf",
            file_path=str(file_path),
            chunks=chunks,
            metadata=metadata,
        )

    def _parse_docx(self, file_path: Path) -> Optional[ParsedDocument]:
        """Парсинг DOCX файла."""
        if not self._has_docx:
            logger.error("python-docx not available")
            return None

        from docx import Document

        doc = Document(file_path)
        chunks = []
        metadata = {}

        # Извлекаем метаданные
        core_props = doc.core_properties
        metadata["title"] = core_props.title or ""
        metadata["author"] = core_props.author or ""
        metadata["created"] = str(core_props.created) if core_props.created else ""

        # Парсим параграфы
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Определяем заголовки
            if para.style.name.startswith("Heading"):
                # Сохраняем предыдущую секцию
                if current_content:
                    chunks.append(
                        DocumentChunk(
                            content="\n".join(current_content),
                            doc_type="docx",
                            section=current_section,
                        )
                    )
                    current_content = []

                current_section = text
            else:
                current_content.append(text)

        # Сохраняем последнюю секцию
        if current_content:
            chunks.append(
                DocumentChunk(
                    content="\n".join(current_content),
                    doc_type="docx",
                    section=current_section,
                )
            )

        # Парсим таблицы
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                chunks.append(
                    DocumentChunk(
                        content=table_text,
                        doc_type="docx",
                        section="[Таблица]",
                    )
                )

        logger.info("DOCX parsed", path=str(file_path), chunks=len(chunks))

        return ParsedDocument(
            filename=file_path.name,
            doc_type="docx",
            file_path=str(file_path),
            chunks=chunks,
            metadata=metadata,
        )

    def _extract_table_text(self, table) -> str:
        """Извлечь текст из таблицы DOCX."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _parse_markdown(self, file_path: Path) -> Optional[ParsedDocument]:
        """Парсинг Markdown/TXT файла."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        chunks = []
        metadata = {"lines": content.count("\n") + 1}

        # Разбиваем по заголовкам
        sections = re.split(r"^(#{1,6}\s+.+)$", content, flags=re.MULTILINE)

        current_section = None
        for i, part in enumerate(sections):
            part = part.strip()
            if not part:
                continue

            if re.match(r"^#{1,6}\s+", part):
                current_section = re.sub(r"^#{1,6}\s+", "", part)
            else:
                chunks.append(
                    DocumentChunk(
                        content=part,
                        doc_type="md",
                        section=current_section,
                    )
                )

        # Если нет заголовков — весь текст как один чанк
        if not chunks and content.strip():
            chunks.append(
                DocumentChunk(
                    content=content.strip(),
                    doc_type="md",
                )
            )

        logger.info("Markdown parsed", path=str(file_path), chunks=len(chunks))

        return ParsedDocument(
            filename=file_path.name,
            doc_type="md",
            file_path=str(file_path),
            chunks=chunks,
            metadata=metadata,
        )

    def _parse_xlsx(self, file_path: Path) -> Optional[ParsedDocument]:
        """Парсинг XLSX файла."""
        if not self._has_openpyxl:
            logger.error("openpyxl not available")
            return None

        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        chunks = []
        metadata = {"sheets": wb.sheetnames}

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_data = []

            for row in ws.iter_rows(values_only=True):
                # Фильтруем пустые строки
                if any(cell is not None for cell in row):
                    row_text = " | ".join(
                        str(cell) if cell is not None else "" for cell in row
                    )
                    rows_data.append(row_text)

            if rows_data:
                chunks.append(
                    DocumentChunk(
                        content="\n".join(rows_data),
                        doc_type="xlsx",
                        section=sheet_name,
                        row_range=f"1-{len(rows_data)}",
                    )
                )

        logger.info("XLSX parsed", path=str(file_path), sheets=len(chunks))

        return ParsedDocument(
            filename=file_path.name,
            doc_type="xlsx",
            file_path=str(file_path),
            chunks=chunks,
            metadata=metadata,
        )


class DocumentLoader:
    """
    Загрузчик документов из папки.

    Сканирует папку и парсит все поддерживаемые файлы.
    """

    def __init__(self):
        """Инициализация загрузчика."""
        self.parser = DocumentParser()
        self._loaded_files: set = set()

    def load_all(self, input_dir: Path) -> List[ParsedDocument]:
        """
        Загрузить все документы из папки.

        Args:
            input_dir: Путь к папке с документами

        Returns:
            Список распарсенных документов
        """
        input_dir = Path(input_dir)

        if not input_dir.exists():
            logger.warning("Input directory not found", path=str(input_dir))
            return []

        documents = []

        # Сканируем все файлы (включая подпапки)
        for file_path in input_dir.rglob("*"):
            if file_path.is_file():
                if file_path.suffix.lower() in DocumentParser.SUPPORTED_EXTENSIONS:
                    doc = self.parser.parse(file_path)
                    if doc:
                        documents.append(doc)
                        self._loaded_files.add(str(file_path))

        logger.info(
            "Documents loaded",
            directory=str(input_dir),
            count=len(documents),
        )

        return documents

    def check_for_new(self, input_dir: Path) -> List[ParsedDocument]:
        """
        Проверить наличие новых документов.

        Args:
            input_dir: Путь к папке с документами

        Returns:
            Список новых документов (ещё не загруженных)
        """
        input_dir = Path(input_dir)

        if not input_dir.exists():
            return []

        new_documents = []

        for file_path in input_dir.rglob("*"):
            if file_path.is_file():
                if str(file_path) not in self._loaded_files:
                    if file_path.suffix.lower() in DocumentParser.SUPPORTED_EXTENSIONS:
                        doc = self.parser.parse(file_path)
                        if doc:
                            new_documents.append(doc)
                            self._loaded_files.add(str(file_path))

        if new_documents:
            logger.info("New documents found", count=len(new_documents))

        return new_documents

    def get_loaded_files(self) -> List[str]:
        """Получить список загруженных файлов."""
        return list(self._loaded_files)
